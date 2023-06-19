import tempfile
import urllib.parse
import urllib.request
import datetime

import pdfkit
import docx2python
from docx import Document
import lxml.etree as etree
from docx.shared import Inches
from bs4 import BeautifulSoup
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session

from core import wkhtmltopdf_path, download_file_to_tempfile


from core import jinja_env

from services import (hr_document_template_service,
                      candidate_service,
                      staff_unit_service,
                      family_relation_service,
                      family_service)
from models import (User,
                    CandidateStageAnswerDefault,
                    CandidateStageAnswerText,
                    CandidateStageType,
                    CandidateStageQuestion,
                    CandidateStageAnswerChoice)
from exceptions import NotFoundException


class RenderService:

    def generate(self, db: Session, candidate_id: str, template_id: str):
        candidate = candidate_service.get_by_id(db, candidate_id)

        candidate_staff_unit = staff_unit_service.get_by_id(
            db, candidate['staff_unit_id'])

        candidate_user: User = candidate_staff_unit.users[0]

        curator_staff_unit = staff_unit_service.get_by_id(
            db, candidate['staff_unit_curator_id'])

        curator_user: User = curator_staff_unit.users[0]

        father_relation = family_relation_service.get_by_name(db, "Отец")
        mother_relation = family_relation_service.get_by_name(db, "Мать")

        father = family_service.get_by_relation_id(db, father_relation.id)
        mother = family_service.get_by_relation_id(db, mother_relation.id)

        document_template = hr_document_template_service.get_by_id(
            db, template_id
        )

        with tempfile.NamedTemporaryFile(delete=False) as temp:
            arr = document_template.path.rsplit(".")
            extension = arr[len(arr) - 1]
            temp_file_path: str = temp.name + "." + extension

            try:
                urllib.request.urlretrieve(
                    document_template.path, temp_file_path)
            except Exception:
                raise NotFoundException(detail="Файл не найден!")

        template = jinja_env.get_template(temp_file_path.replace('/tmp/', ''))

        now = datetime.datetime.now()

        month_num = now.month

        month_names = [
            'Қаңтар', 'Ақпан', 'Наурыз',
            'Сәуір', 'Мамыр', 'Маусым',
            'Шілде', 'Тамыз', 'Қыркүйек',
            'Қазан', 'Қараша', 'Желтоқсан'
        ]

        month_name = month_names[month_num - 1]

        context = {
            'candidate': candidate_user,
            'curator': curator_user,
            'year': now.year,
            'month': month_name,
            'day': now.day,
            'father': father,
            'mother': mother
        }

        ans = template.render(context)

        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            file_name = temp_file.name + "." + extension
            with open(file_name, "w") as f:
                f.write(ans)

        return FileResponse(
            path=file_name,
            filename=document_template.name + "." + extension,
        )

    async def generate_finish_candidate(self,
                                        db: Session,
                                        candidate_id: str,
                                        template_id: str):
        hr_document_template = hr_document_template_service.get_by_id(
            db, template_id)
        candidate = candidate_service.get_by_id(db, candidate_id)

        temp_file_path = await download_file_to_tempfile(hr_document_template.pathKZ)
        template = jinja_env.get_template(temp_file_path.replace('/tmp/', ''))

        candidate_staff_unit = staff_unit_service.get_by_id(
            db, candidate.staff_unit_id)
        user = candidate_staff_unit.users[0]

        # autobiography
        autobigoraphy_question = db.query(CandidateStageQuestion).filter(
            CandidateStageQuestion.question == "Краткие сведения из автобиографии"
        ).first()

        autobiography_answer = db.query(CandidateStageAnswerDefault).filter(
            CandidateStageAnswerDefault.candidate_id == candidate_id,
            CandidateStageAnswerDefault.candidate_stage_question_id
            == autobigoraphy_question.id
        ).first()

        # essay
        essay_stage_type = (db.query(CandidateStageType)
                            .filter(CandidateStageType.name == 'Рецензия на эссе')
                            .first())

        essay_stage_question = db.query(CandidateStageQuestion).filter(
            CandidateStageQuestion.candidate_stage_type_id == essay_stage_type.id
        ).first()

        essay_answer = db.query(CandidateStageAnswerText).filter(
            CandidateStageAnswerText.candidate_id == candidate_id,
            CandidateStageAnswerText.candidate_stage_question_id
            == essay_stage_question.id
        ).first()

        # military medical commision
        military_medical_commision_stage = (db.query(CandidateStageType)
                                            .filter(CandidateStageType.name
                                                    == 'Военно-врачебная комиссия')
                                            .first())

        military_medical_commision_question = db.query(CandidateStageQuestion).filter(
            CandidateStageQuestion.candidate_stage_type_id
            == military_medical_commision_stage.id
        ).first()

        military_medical_commision_answer = db.query(CandidateStageAnswerChoice).filter(
            CandidateStageAnswerText.candidate_id == candidate_id,
            CandidateStageAnswerText.candidate_stage_question_id
            == military_medical_commision_question.id
        ).first()

        # psychological stage
        psychological_stage_type = (db.query(CandidateStageType)
                                    .filter(CandidateStageType.name
                                            == 'Беседа с психологом')
                                    .first())

        psychological_stage_question = db.query(CandidateStageQuestion).filter(
            CandidateStageQuestion.candidate_stage_type_id
            == psychological_stage_type.id
        ).first()

        psychological_stage_answer = db.query(CandidateStageAnswerText).filter(
            CandidateStageAnswerText.candidate_id == candidate_id,
            CandidateStageAnswerText.candidate_stage_question_id
            == psychological_stage_question.id
        ).first()

        # polygraph stage
        polygraph_stage_type = (db
                    .query(CandidateStageType)
                    .filter(CandidateStageType.name
                    == 'Результаты полиграфологического исследования')
                    .first())

        polygraph_stage_question = db.query(CandidateStageQuestion).filter(
            CandidateStageQuestion.candidate_stage_type_id == polygraph_stage_type.id
        ).first()

        polygraph_stage_answer = db.query(CandidateStageAnswerChoice).filter(
            CandidateStageAnswerChoice.candidate_id == candidate_id,
            CandidateStageAnswerChoice.candidate_stage_question_id
            == polygraph_stage_question.id
        ).first()

        # sport stage
        sport_stage_type = (db.query(CandidateStageType)
                            .filter(CandidateStageType.name
                                    == 'Результаты физической подготовки')
                            .first())

        sport_stage_question = db.query(CandidateStageQuestion).filter(
            CandidateStageQuestion.candidate_stage_type_id == sport_stage_type.id
        ).first()

        sport_stage_answer = db.query(CandidateStageAnswerChoice).filter(
            CandidateStageAnswerChoice.candidate_id == candidate_id,
            CandidateStageAnswerChoice.candidate_stage_question_id
            == sport_stage_question.id
        ).first()

        # legislation stage
        legislation_stage_type = (db
                    .query(CandidateStageType)
                    .filter(CandidateStageType.name
                    == 'Результаты тестирования на знание законодательства РК')
                    .first())

        legislation_stage_question = db.query(CandidateStageQuestion).filter(
            CandidateStageQuestion.candidate_stage_type_id == legislation_stage_type.id
        ).first()

        legislation_stage_answer = db.query(CandidateStageAnswerChoice).filter(
            CandidateStageAnswerChoice.candidate_id == candidate_id,
            CandidateStageAnswerChoice.candidate_stage_question_id
            == legislation_stage_question.id
        ).first()

        # Wife
        wife_relation = family_relation_service.get_by_name(db, "Жена")
        wife = family_service.get_by_relation_id(db, wife_relation.id)

        # Husband
        husband_relation = family_relation_service.get_by_name(db, "Муж")
        husband = family_service.get_by_relation_id(db, husband_relation.id)

        # Father
        father_relation = family_relation_service.get_by_name(db, "Отец")
        father = family_service.get_by_relation_id(db, father_relation.id)

        # Mother
        mother_relation = family_relation_service.get_by_name(db, "Мать")
        mother = family_service.get_by_relation_id(db, mother_relation.id)

        # Brother
        brother_relation = family_relation_service.get_by_name(db, "Брат")
        brother = family_service.get_by_relation_id(db, brother_relation.id)

        context = {
            'user': user,
            'candidate': candidate,
            'autobiography_answer': autobiography_answer,
            'essay_answer': essay_answer,
            'military_medical_commision_answer': military_medical_commision_answer,
            'psychological_stage_answer': psychological_stage_answer,
            'polygraph_stage_answer': polygraph_stage_answer,
            'sport_stage_answer': sport_stage_answer,
            'legislation_stage_answer': legislation_stage_answer,
            'department_planned': None,
            'wife': wife,
            'husband': husband,
            'father': father,
            'mother': mother,
            'brother': brother
        }

        ans = template.render(context)

        opts = {
            'encoding': 'UTF-8',
            'enable-local-file-access': True
        }

        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            file_name = temp_file.name + ".pdf"
            pdfkit.from_string(ans,
                               file_name,
                               options=opts,
                               configuration=pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path))

        return FileResponse(
            path=file_name,
            filename=(
                f"Заключение о зачислении кандидата {user.last_name} "
                f"{user.first_name}.pdf"
            ),
        )

    def convert_html_to_docx(self, html: str):
        soup = BeautifulSoup(html, "html.parser")

        document = Document()

        for tag in soup.recursiveChildGenerator():
            if tag.name == 'p':
                document.add_paragraph(tag.text)
            elif tag.name == 'img':
                document.add_picture(tag['src'], width=Inches(2))

        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            file_name = temp_file.name

            document.save(file_name)

        return FileResponse(
            path=file_name,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="test.docx",
        )

    def convert_docx_to_xml_to_html(self):
        doc_result = docx2python.docx2python('finish.docx')
        xml = doc_result.to_xml()

        # Parse the XML file
        root = etree.fromstring(xml)

        # Add the namespace for WordprocessingML
        etree.register_namespace(
            'w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')

        # Add the namespace for Office Art
        etree.register_namespace(
            'a', 'http://schemas.openxmlformats.org/drawingml/2006/main')

        # Add the namespace for DrawingML
        etree.register_namespace(
            'wp', 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing')

        # Create an XSLT stylesheet to transform the XML to HTML
        xslt = etree.XML(
            '''
        <xsl:stylesheet
            version="1.0"
            xmlns:xsl="http://www.w3.org/1999/XSL/Transform"
            xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
            xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
        <xsl:output method="html"/>

        <!-- Map Word styles to HTML styles -->
        <xsl:template match="w:p">
            <p>
            <xsl:apply-templates/>
            </p>
        </xsl:template>

        <xsl:template match="w:r">
            <span>
            <xsl:apply-templates/>
            </span>
        </xsl:template>

        <xsl:template match="w:t">
            <xsl:value-of select="."/>
        </xsl:template>

        <xsl:template match="w:pStyle">
            <xsl:attribute name="class">
            <xsl:value-of select="."/>
            </xsl:attribute>
        </xsl:template>

        <xsl:template match="w:rPr/w:b">
            <xsl:attribute name="style">
            <xsl:text>font-weight:bold;</xsl:text>
            </xsl:attribute>
        </xsl:template>

        <xsl:template match="w:rPr/w:i">
            <xsl:attribute name="style">
            <xsl:text>font-style:italic;</xsl:text>
            </xsl:attribute>
        </xsl:template>

        <xsl:template match="w:rPr/w:u">
            <xsl:attribute name="style">
            <xsl:text>text-decoration:underline;</xsl:text>
            </xsl:attribute>
        </xsl:template>

        <xsl:template match="w:pPr/w:jc">
            <xsl:attribute name="style">
            <xsl:text>text-align:</xsl:text>
            <xsl:value-of select="@w:val"/>
            <xsl:text>;</xsl:text>
            </xsl:attribute>
        </xsl:template>

        </xsl:stylesheet>''')

        # Transform the XML to HTML using the XSLT stylesheet
        transform = etree.XSLT(xslt)
        result = transform(root)

        # Save the resulting HTML to a file
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            file_name = temp_file.name + ".html"
            with open(file_name, 'wb') as f:
                f.write(result)

        return FileResponse(
            path=file_name,
        )

    def convert_html_to_pdf(self, html: str):
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            file_name = temp_file.name + ".pdf"
            options = {
                'encoding': 'UTF-8',
                'enable-local-file-access': True
            }
            pdfkit.from_string(html,
                               file_name,
                               options=options,
                               configuration=pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path))

        return FileResponse(
            path=file_name,
            filename='result.pdf',
        )


render_service = RenderService()
